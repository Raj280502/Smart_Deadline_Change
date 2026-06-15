import os
import re
from pathlib import Path
from urllib.parse import urlparse

import httpx


DOWNLOAD_DIR = Path("storage") / "placement_documents"


def download_document(url: str, filename_hint: str = "jd") -> str:
    """
    Download a JD file from the portal.

    Stack syntax:
        httpx.get(url)        -> fetches the file
        Path.write_bytes(...) -> saves bytes to disk
    """
    DOWNLOAD_DIR.mkdir(parents=True, exist_ok=True)

    response = httpx.get(url, timeout=30, follow_redirects=True)
    response.raise_for_status()

    extension = _extension_from_url(url) or ".pdf"
    filename = f"{_safe_name(filename_hint)}{extension}"
    path = DOWNLOAD_DIR / filename
    path.write_bytes(response.content)
    return str(path)


def extract_text_from_document(path: str) -> str:
    """Extract text from PDF, DOCX, or TXT job-description files."""
    suffix = Path(path).suffix.lower()

    if suffix == ".pdf":
        return extract_text_from_pdf(path)
    if suffix == ".docx":
        return extract_text_from_docx(path)
    if suffix in (".txt", ".text"):
        return Path(path).read_text(encoding="utf-8", errors="ignore")

    return ""


def extract_text_from_pdf(path: str) -> str:
    """
    Read PDF text with pypdf.

    Stack syntax:
        reader = PdfReader(path)
        page.extract_text()
    """
    from pypdf import PdfReader

    reader = PdfReader(path)
    parts = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            parts.append(text.strip())
    return _clean_text("\n".join(parts))


def extract_text_from_docx(path: str) -> str:
    """
    Read Word JD text with python-docx.

    Stack syntax:
        Document(path).paragraphs
    """
    from docx import Document

    document = Document(path)
    text = "\n".join(p.text for p in document.paragraphs if p.text.strip())
    return _clean_text(text)


def _extension_from_url(url: str) -> str:
    path = urlparse(url).path
    extension = os.path.splitext(path)[1].lower()
    return extension if extension in (".pdf", ".docx", ".txt") else ""


def _safe_name(value: str) -> str:
    cleaned = re.sub(r"[^a-zA-Z0-9_-]+", "_", value).strip("_")
    return cleaned[:80] or "jd"


def _clean_text(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()
